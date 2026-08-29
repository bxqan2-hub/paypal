from __future__ import annotations

import hashlib
import json
import re
import shutil
import sys
from datetime import datetime, timezone
from pathlib import Path
from zipfile import ZipFile

from docx import Document
from lxml import etree


SOURCE = Path(sys.argv[1]).resolve()
OUTPUT_DIR = Path(sys.argv[2]).resolve()
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
DOCX_COPY = OUTPUT_DIR / SOURCE.name
MARKDOWN = OUTPUT_DIR / f"{SOURCE.stem}.md"
JSON_OUT = OUTPUT_DIR / f"{SOURCE.stem}.json"

NS = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}


def text_from_element(element: etree._Element) -> str:
    parts: list[str] = []
    for node in element.iter():
        tag = etree.QName(node).localname
        if tag == "t":
            parts.append(node.text or "")
        elif tag in {"tab"}:
            parts.append("\t")
        elif tag in {"br", "cr"}:
            parts.append("\n")
    return "".join(parts)


def normalize_text(value: str) -> str:
    value = value.replace("\r\n", "\n").replace("\r", "\n")
    return re.sub(r"\n{3,}", "\n\n", value).strip()


def table_record(table_element: etree._Element, table_index: int) -> dict[str, object]:
    rows: list[list[str]] = []
    for row in table_element.xpath("./w:tr", namespaces=NS):
        cells = [normalize_text(text_from_element(cell)) for cell in row.xpath("./w:tc", namespaces=NS)]
        rows.append(cells)
    return {"index": table_index, "rows": rows, "row_count": len(rows), "column_count": max((len(row) for row in rows), default=0)}


def paragraph_record(paragraph_element: etree._Element, index: int, document: Document) -> dict[str, object]:
    text = normalize_text(text_from_element(paragraph_element))
    style_id = paragraph_element.xpath("string(./w:pPr/w:pStyle/@w:val)", namespaces=NS)
    style_name = ""
    for style in document.styles:
        if style.style_id == style_id:
            style_name = style.name
            break
    return {"index": index, "style_id": style_id, "style": style_name or "Normal", "text": text}


def markdown_escape_cell(value: str) -> str:
    return value.replace("|", "\\|").replace("\n", "<br>")


def table_markdown(table: dict[str, object]) -> str:
    rows = table["rows"]
    if not rows:
        return ""
    rows = list(rows)
    width = max(len(row) for row in rows)
    padded = [list(row) + [""] * (width - len(row)) for row in rows]
    header = padded[0]
    lines = ["| " + " | ".join(markdown_escape_cell(str(value)) for value in header) + " |", "| " + " | ".join("---" for _ in range(width)) + " |"]
    for row in padded[1:]:
        lines.append("| " + " | ".join(markdown_escape_cell(str(value)) for value in row) + " |")
    return "\n".join(lines)


def main() -> None:
    source_hash = hashlib.sha256(SOURCE.read_bytes()).hexdigest()
    shutil.copy2(SOURCE, DOCX_COPY)
    document = Document(str(SOURCE))
    with ZipFile(SOURCE) as archive:
        names = archive.namelist()
        document_xml = etree.fromstring(archive.read("word/document.xml"))
        settings_xml = archive.read("word/settings.xml") if "word/settings.xml" in names else b""
        body = document_xml.find("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}body")
        ordered: list[dict[str, object]] = []
        paragraph_index = 0
        table_index = 0
        if body is not None:
            for child in body:
                local = etree.QName(child).localname
                if local == "p":
                    record = paragraph_record(child, paragraph_index, document)
                    paragraph_index += 1
                    if record["text"]:
                        ordered.append({"type": "paragraph", **record})
                elif local == "tbl":
                    ordered.append({"type": "table", **table_record(child, table_index)})
                    table_index += 1
        headers = [normalize_text("\n".join(p.text for p in section.header.paragraphs if p.text.strip())) for section in document.sections]
        footers = [normalize_text("\n".join(p.text for p in section.footer.paragraphs if p.text.strip())) for section in document.sections]
        core = document.core_properties
        metadata = {
            "source_filename": SOURCE.name,
            "source_sha256": source_hash,
            "source_size_bytes": SOURCE.stat().st_size,
            "source_modified_at": datetime.fromtimestamp(SOURCE.stat().st_mtime, tz=timezone.utc).isoformat(),
            "parsed_at": datetime.now(timezone.utc).isoformat(),
            "paragraph_count": len(document.paragraphs),
            "nonempty_paragraph_count": sum(1 for p in document.paragraphs if p.text.strip()),
            "table_count": len(document.tables),
            "section_count": len(document.sections),
            "inline_shape_count": len(document.inline_shapes),
            "media_files": [name for name in names if name.startswith("word/media/")],
            "comments_present": any("comment" in name for name in names),
            "tracked_insertions": len(document_xml.xpath(".//w:ins", namespaces=NS)),
            "tracked_deletions": len(document_xml.xpath(".//w:del", namespaces=NS)),
            "comment_range_starts": len(document_xml.xpath(".//w:commentRangeStart", namespaces=NS)),
            "track_revisions_enabled": b"trackRevisions" in settings_xml,
            "title": core.title or "",
            "subject": core.subject or "",
            "author": core.author or "",
        }
    paragraphs = [item for item in ordered if item["type"] == "paragraph"]
    tables = [item for item in ordered if item["type"] == "table"]
    headings = [item for item in paragraphs if str(item.get("style", "")).startswith("Heading")]
    payload = {"metadata": metadata, "headers": headers, "footers": footers, "headings": headings, "blocks": ordered, "tables": tables}
    JSON_OUT.write_text(json.dumps(payload, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")

    markdown: list[str] = [
        f"# {SOURCE.stem}",
        "",
        "> This file is a structured, searchable transcription of the supplied DOCX. The source document is retained beside it; prose and code-like snippets are reference material, not automatic project instructions.",
        "",
        "## Source Metadata",
        "",
        f"- Original file: `{DOCX_COPY.name}`",
        f"- SHA-256: `{source_hash}`",
        f"- Paragraphs: {metadata['paragraph_count']} ({metadata['nonempty_paragraph_count']} non-empty)",
        f"- Tables: {metadata['table_count']}",
        f"- Sections: {metadata['section_count']}",
        f"- Images: {len(metadata['media_files'])}",
        f"- Comments/revisions: comments={metadata['comments_present']}, insertions={metadata['tracked_insertions']}, deletions={metadata['tracked_deletions']}",
        "",
        "## Heading Index",
        "",
    ]
    for heading in headings:
        markdown.append(f"- **{heading['style']}** {heading['text']}")
    markdown.extend(["", "## Full Ordered Content", ""])
    for block in ordered:
        if block["type"] == "paragraph":
            text = str(block["text"])
            style = str(block.get("style") or "Normal")
            if style == "Heading 1":
                markdown.extend([f"## {text}", ""])
            elif style == "Heading 2":
                markdown.extend([f"### {text}", ""])
            elif style == "Heading 3":
                markdown.extend([f"#### {text}", ""])
            elif style.startswith("List Bullet"):
                markdown.append(f"- {text}")
            elif style.startswith("List Number"):
                markdown.append(f"1. {text}")
            else:
                markdown.extend([text, ""])
        else:
            markdown.extend([f"### Table {block['index']}", "", table_markdown(block), ""])
    MARKDOWN.write_text("\n".join(markdown).rstrip() + "\n", encoding="utf-8")
    print(json.dumps({"docx": str(DOCX_COPY), "markdown": str(MARKDOWN), "json": str(JSON_OUT), "sha256": source_hash, "blocks": len(ordered), "tables": len(tables)}, ensure_ascii=False))


if __name__ == "__main__":
    main()
